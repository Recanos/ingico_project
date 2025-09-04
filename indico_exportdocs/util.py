from docx import Document
from indico.modules.events.models.events import Event
from indico.modules.events.papers.models.revisions import PaperRevisionState
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor

def generate_docx_list(event_id):
    event = Event.get(event_id)
    doc = Document()
    
    def set_black_color(element):
        if hasattr(element, 'runs'):
            for run in element.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        if hasattr(element, 'paragraphs'):
            for paragraph in element.paragraphs:
                set_black_color(paragraph)
        if hasattr(element, 'tables'):
            for table in element.tables:
                set_black_color(table)
    
    # Настройка полей страницы по ГОСТ Р 7.0.97-2016
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.79)    
        section.right_margin = Inches(0.39)  
        section.top_margin = Inches(0.79)     
        section.bottom_margin = Inches(0.79)  
    
    # Добавляем заголовок документа
    title = doc.add_heading(f'СПИСОК ДОКЛАДОВ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Добавляем подзаголовок с названием события
    subtitle = doc.add_paragraph(f'"{event.title}"')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0) 
    
    doc.add_paragraph()
    
    contributions = [c for c in event.contributions if not c.is_deleted and c.start_dt]
    
    from collections import defaultdict
    
    # Группируем по дате
    date_groups = defaultdict(list)
    for contrib in contributions:
        # Получаем дату без времени
        date_key = contrib.start_dt.date()
        date_groups[date_key].append(contrib)
    
    # Если есть доклады с временем, группируем по датам
    if date_groups:
        # Сортируем даты
        sorted_dates = sorted(date_groups.keys())
        
        for date_key in sorted_dates:
            date_contributions = date_groups[date_key]
            
            # Сортируем доклады по времени в рамках дня
            date_contributions.sort(key=lambda x: x.start_dt)
            
            date_str = date_key.strftime('%d %B %Y г.')

            month_translations = {
                'January': 'января', 'February': 'февраля', 'March': 'марта', 'April': 'апреля',
                'May': 'мая', 'June': 'июня', 'July': 'июля', 'August': 'августа',
                'September': 'сентября', 'October': 'октября', 'November': 'ноября', 'December': 'декабря'
            }
            for eng, rus in month_translations.items():
                date_str = date_str.replace(eng, rus)
            
            date_heading = doc.add_heading(f'Заседание {len(sorted_dates) if len(sorted_dates) > 1 else ""}', level=1)
            date_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in date_heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            date_para = doc.add_paragraph(f'{date_str}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'  
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Фамилия и инициалы докладчика, название доклада'
            hdr_cells[2].text = 'Статус (магистр / студент)'
            hdr_cells[3].text = 'Решение'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            row_number = 1
            for contribution in date_contributions:
                speakers = [link.person for link in contribution.person_links if link.is_speaker]
                
                if not speakers:
                    continue
                    
                for speaker in speakers:
                    row = table.add_row().cells
                    
                    row[0].text = str(row_number)
                    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    speaker_name = f"{speaker.last_name} {speaker.first_name[0]}.{speaker.first_name[1] if len(speaker.first_name) > 1 else ''}"
                    contribution_title = contribution.title or 'Без названия'
                    row[1].text = f"{speaker_name}. {contribution_title}"
                    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Колонка "Статус (магистр / студент)"
                    status = determine_student_status(speaker)
                    row[2].text = status
                    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    row[3].text = ''
                    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    row_number += 1
            
            # Добавляем пустую строку между датами
            doc.add_paragraph()
    
    # Доклады без времени
    no_time_contribs = [c for c in event.contributions if not c.is_deleted and not c.start_dt]
    if no_time_contribs:
        no_time_heading = doc.add_heading('Доклады без указанного времени', level=1)
        for run in no_time_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Создаем таблицу для докладов без времени
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Настройка заголовков таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Фамилия и инициалы докладчика, название доклада'
        hdr_cells[2].text = 'Статус (магистр / студент)'
        hdr_cells[3].text = 'Решение'
        
        # Форматирование заголовков таблицы
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Сортируем доклады по названию для удобства
        no_time_contribs.sort(key=lambda x: x.title.lower() if x.title else '')
        
        row_number = 1
        for contribution in no_time_contribs:
            # Получаем докладчиков (speakers)
            speakers = [link.person for link in contribution.person_links if link.is_speaker]
            
            if not speakers:
                continue
                
            for speaker in speakers:
                row = table.add_row().cells
                
                # Колонка № - порядковый номер
                row[0].text = str(row_number)
                row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Колонка "Фамилия и инициалы докладчика, название доклада"
                speaker_name = f"{speaker.last_name} {speaker.first_name[0]}.{speaker.first_name[1] if len(speaker.first_name) > 1 else ''}"
                contribution_title = contribution.title or 'Без названия'
                row[1].text = f"{speaker_name}. {contribution_title}"
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Колонка "Статус (магистр / студент)"
                status = determine_student_status(speaker)
                row[2].text = status
                row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Колонка "Решение" - оставляем пустой для ручного заполнения
                row[3].text = ''
                row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                row_number += 1
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0) 
    
    # Настройка межстрочного интервала
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
    
    # Настройка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(0, 0, 0)  
                    paragraph.paragraph_format.line_spacing = 1.5
    
    set_black_color(doc)
    
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def determine_student_status(person):
    """
    Определяет статус участника (студент/магистр) на основе affiliation и других данных
    """
    if not person.affiliation:
        return 'Не указан'
    
    affiliation_lower = person.affiliation.lower()
    
    # Ключевые слова для определения статуса
    student_keywords = ['студент', 'student', 'бакалавр', 'bachelor', '1 курс', '2 курс', '3 курс', '4 курс']
    master_keywords = ['магистр', 'master', 'магистрант', '5 курс', '6 курс']
    
    for keyword in student_keywords:
        if keyword in affiliation_lower:
            return 'Студент'
    
    for keyword in master_keywords:
        if keyword in affiliation_lower:
            return 'Магистр'
    
    return person.affiliation

def generate_docx_report(event_id):
    event = Event.get(event_id)
    doc = Document()
    
    def set_black_color(element):
        if hasattr(element, 'runs'):
            for run in element.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        if hasattr(element, 'paragraphs'):
            for paragraph in element.paragraphs:
                set_black_color(paragraph)
        if hasattr(element, 'tables'):
            for table in element.tables:
                set_black_color(table)
    
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.79) 
        section.right_margin = Inches(0.39)
        section.top_margin = Inches(0.79)     
        section.bottom_margin = Inches(0.79)  
    
    title = doc.add_heading(f'ОТЧЕТ О ПРОВЕДЕНИИ КОНФЕРЕНЦИИ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Добавляем подзаголовок с названием события
    subtitle = doc.add_paragraph(f'"{event.title}"')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0) 
    
    doc.add_paragraph()
    
    contributions = [c for c in event.contributions if not c.is_deleted and c.start_dt]
    
    from collections import defaultdict
    
    date_groups = defaultdict(list)
    for contrib in contributions:
        # Получаем дату без времени
        date_key = contrib.start_dt.date()
        date_groups[date_key].append(contrib)

    if date_groups:
        # Сортируем даты
        sorted_dates = sorted(date_groups.keys())
        
        for date_key in sorted_dates:
            date_contributions = date_groups[date_key]
            
            date_contributions.sort(key=lambda x: x.start_dt)
            
            date_str = date_key.strftime('%d %B %Y г., %H-%M')
            month_translations = {
                'January': 'января', 'February': 'февраля', 'March': 'марта', 'April': 'апреля',
                'May': 'мая', 'June': 'июня', 'July': 'июля', 'August': 'августа',
                'September': 'сентября', 'October': 'октября', 'November': 'ноября', 'December': 'декабря'
            }
            for eng, rus in month_translations.items():
                date_str = date_str.replace(eng, rus)
            
            date_heading = doc.add_heading(f'Заседание {len(sorted_dates) if len(sorted_dates) > 1 else ""}', level=1)
            date_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in date_heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            date_para = doc.add_paragraph(f'{date_str}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            row_number = 1
            for contribution in date_contributions:

                speakers = [link.person for link in contribution.person_links if link.is_speaker]
                
                if not speakers:
                    continue
                    
                for speaker in speakers:
                    p = doc.add_paragraph()
                    
                    run_number = p.add_run(f"{row_number}. ")
                    run_number.bold = True
                    
                    speaker_name = f"{speaker.last_name} {speaker.first_name[0]}.{speaker.first_name[1] if len(speaker.first_name) > 1 else ''}"
                    run_name = p.add_run(speaker_name)
                    run_name.bold = True
                    
                    contribution_title = contribution.title or 'Без названия'
                    p.add_run(f". {contribution_title}")
                    
                    row_number += 1
            
            doc.add_paragraph()
    
    # Доклады без времени
    no_time_contribs = [c for c in event.contributions if not c.is_deleted and not c.start_dt]
    if no_time_contribs:
        no_time_heading = doc.add_heading('Доклады без указанного времени', level=1)
        for run in no_time_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        no_time_contribs.sort(key=lambda x: x.title.lower() if x.title else '')
        
        row_number = 1
        for contribution in no_time_contribs:
            speakers = [link.person for link in contribution.person_links if link.is_speaker]
            
            if not speakers:
                continue
                
            for speaker in speakers:
                p = doc.add_paragraph()
                
                run_number = p.add_run(f"{row_number}. ")
                run_number.bold = True
                
                speaker_name = f"{speaker.last_name} {speaker.first_name[0]}.{speaker.first_name[1] if len(speaker.first_name) > 1 else ''}"
                run_name = p.add_run(speaker_name)
                run_name.bold = True
                
                contribution_title = contribution.title or 'Без названия'
                p.add_run(f". {contribution_title}")
                
                row_number += 1
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0) 
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
    
    set_black_color(doc)
    
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def generate_docx_papers(event_id):
    event = Event.get(event_id)
    doc = Document()
    
    def set_black_color(element):
        if hasattr(element, 'runs'):
            for run in element.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        if hasattr(element, 'paragraphs'):
            for paragraph in element.paragraphs:
                set_black_color(paragraph)
        if hasattr(element, 'tables'):
            for table in element.tables:
                set_black_color(table)
    
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.79)    
        section.right_margin = Inches(0.39)  
        section.top_margin = Inches(0.79)     
        section.bottom_margin = Inches(0.79)  
    
    title = doc.add_heading(f'СПИСОК ПУБЛИКАЦИЙ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    subtitle = doc.add_paragraph(f'"{event.title}"')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0) 
    
    doc.add_paragraph()
    
    contributions = [c for c in event.contributions if not c.is_deleted and c.start_dt]
    
    from collections import defaultdict
    
    date_groups = defaultdict(list)
    for contrib in contributions:
        date_key = contrib.start_dt.date()
        date_groups[date_key].append(contrib)
    
    if date_groups:
        sorted_dates = sorted(date_groups.keys())
        
        for date_key in sorted_dates:
            date_contributions = date_groups[date_key]
            
            date_contributions.sort(key=lambda x: x.start_dt)
            
            date_str = date_key.strftime('%d %B, %H:%M')
            month_translations = {
                'January': 'января', 'February': 'февраля', 'March': 'марта', 'April': 'апреля',
                'May': 'мая', 'June': 'июня', 'July': 'июля', 'August': 'августа',
                'September': 'сентября', 'October': 'октября', 'November': 'ноября', 'December': 'декабря'
            }
            for eng, rus in month_translations.items():
                date_str = date_str.replace(eng, rus)
            
            date_heading = doc.add_heading(f'Заседание {len(sorted_dates) if len(sorted_dates) > 1 else ""}.', level=1)
            date_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in date_heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            date_para = doc.add_paragraph(f'{date_str}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            row_number = 1
            for contribution in date_contributions:
                if hasattr(contribution, '_accepted_paper_revision') and contribution._accepted_paper_revision:
                    rev = contribution._accepted_paper_revision
                    if hasattr(rev, 'state') and rev.state == PaperRevisionState.accepted:
                        authors = [link.person for link in contribution.person_links if link.is_speaker]
                        
                        if not authors:
                            continue
                            
                        for author in authors:
                            p = doc.add_paragraph()
                            
                            run_number = p.add_run(f"    {row_number}. ")
                            run_number.bold = True
                            
                            full_name = f"{author.first_name} {author.last_name}"
                            if hasattr(author, 'middle_name') and author.middle_name:
                                full_name = f"{author.last_name} {author.first_name} {author.middle_name}"
                            
                            run_name = p.add_run(full_name)
                            run_name.bold = True
                            
                            # Добавляем группу (из affiliation)
                            if author.affiliation:
                                p.add_run(f", {author.affiliation}")
                            
                            p.add_run("\n")
                            
                            article_title = contribution.title or 'Без названия'
                            p.add_run(article_title)
                            
                            row_number += 1
            
            doc.add_paragraph()
    
    no_time_contribs = [c for c in event.contributions if not c.is_deleted and not c.start_dt]
    if no_time_contribs:
        no_time_heading = doc.add_heading('Доклады без указанного времени', level=1)
        for run in no_time_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Получаем принятые к публикации статьи без времени
        row_number = 1
        for contribution in no_time_contribs:
            if hasattr(contribution, '_accepted_paper_revision') and contribution._accepted_paper_revision:
                rev = contribution._accepted_paper_revision
                # Проверяем статус - только accepted
                if hasattr(rev, 'state') and rev.state == PaperRevisionState.accepted:
                    authors = [link.person for link in contribution.person_links if link.is_speaker]
                    
                    if not authors:
                        continue
                        
                    for author in authors:
                        p = doc.add_paragraph()
                        
                        run_number = p.add_run(f"    {row_number}. ")
                        run_number.bold = True
                        
                        full_name = f"{author.first_name} {author.last_name}"
                        if hasattr(author, 'middle_name') and author.middle_name:
                            full_name = f"{author.last_name} {author.first_name} {author.middle_name}"
                        
                        run_name = p.add_run(full_name)
                        run_name.bold = True
                        
                        # Добавляем группу (из affiliation)
                        if author.affiliation:
                            p.add_run(f", {author.affiliation}")
                        
                        p.add_run("\n")
                        
                        article_title = contribution.title or 'Без названия'
                        p.add_run(article_title)
                        
                        row_number += 1
    
    if row_number == 1:
        p = doc.add_paragraph()
        p.add_run("Статьи, принятые к публикации, не найдены.")
        p.italic = True
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0) 
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
    
    set_black_color(doc)
    
    f = BytesIO()
    doc.save(f)
    return f.getvalue()
